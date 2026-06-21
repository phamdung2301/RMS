import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_shading(cell, color_hex):
    """Applies a background fill color to a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Sets padding inside a table cell in dxa (1 pt = 20 dxa)."""
    tcMar_xml = f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>'
    cell._tc.get_or_add_tcPr().append(parse_xml(tcMar_xml))

def set_table_borders(table):
    """Applies clean, modern horizontal-only borders (no vertical lines) to a table."""
    tblPr = table._tbl.tblPr
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="1E3A8A"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="1E3A8A"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))

def add_page_number(run):
    """Applies Word XML fields to insert automated page numbering in footer."""
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def parse_and_add_runs(paragraph, text):
    """Parses inline markdown formatters (**, *, `) and links, appending formatted runs."""
    text = text.replace("<br>", "\n").replace("<br />", "\n")
    
    # Matches: ***bold-italic***, **bold**, *italic*, `code`, and http/https urls
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|https?://\S+)')
    
    parts = []
    last_idx = 0
    
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_idx:
            parts.append((text[last_idx:start], "normal"))
            
        val = match.group(0)
        if val.startswith("***") and val.endswith("***"):
            parts.append((val[3:-3], "bold_italic"))
        elif val.startswith("**") and val.endswith("**"):
            parts.append((val[2:-2], "bold"))
        elif val.startswith("*") and val.endswith("*"):
            parts.append((val[1:-1], "italic"))
        elif val.startswith("`") and val.endswith("`"):
            parts.append((val[1:-1], "code"))
        elif val.startswith("http"):
            parts.append((val, "link"))
        last_idx = end
        
    if last_idx < len(text):
        parts.append((text[last_idx:], "normal"))
        
    for text_part, style in parts:
        if not text_part:
            continue
        run = paragraph.add_run(text_part)
        run.font.name = 'Calibri'
        
        if style == "bold":
            run.bold = True
            run.font.color.rgb = RGBColor(17, 24, 39) # Dark gray
        elif style == "italic":
            run.italic = True
        elif style == "bold_italic":
            run.bold = True
            run.italic = True
        elif style == "code":
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(180, 83, 9) # Warm Amber
        elif style == "link":
            run.underline = True
            run.font.color.rgb = RGBColor(37, 99, 235) # Royal Blue

def compile_markdown_to_docx(md_path, docx_path):
    print(f"Compiling {os.path.basename(md_path)} to {os.path.basename(docx_path)}...")
    
    doc = docx.Document()
    
    # Configure 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Configure footer with automated page number
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Trang ")
        f_run.font.name = "Calibri"
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(107, 114, 128)
        add_page_number(f_p.add_run())
    
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    code_block_lines = []
    
    in_table = False
    table_rows = []
    
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        
        # 1. Handle Code Block
        if stripped.startswith("```"):
            if in_code_block:
                # Closing code block, write accumulated code
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                # Apply background shading to the entire block paragraph
                pBdr_xml = f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="1E3A8A"/></w:pBdr>'
                shading_xml = f'<w:shd {nsdecls("w")} w:fill="F3F4F6"/>'
                p._p.get_or_add_pPr().append(parse_xml(pBdr_xml))
                p._p.get_or_add_pPr().append(parse_xml(shading_xml))
                
                code_text = "".join(code_block_lines)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(55, 65, 81)
                
                code_block_lines = []
            else:
                # Opening code block
                in_code_block = True
            idx += 1
            continue
            
        if in_code_block:
            code_block_lines.append(line)
            idx += 1
            continue
            
        # 2. Handle Markdown Table
        if stripped.startswith("|"):
            in_table = True
            table_rows.append(stripped)
            idx += 1
            continue
        elif in_table:
            # Table ended, compile it
            in_table = False
            # Parse table rows
            parsed_rows = []
            for r in table_rows:
                # Split by | and strip
                cells = [c.strip() for c in r.split("|")[1:-1]]
                parsed_rows.append(cells)
                
            # Filter out separator row (contains ---)
            data_rows = []
            for row in parsed_rows:
                if len(row) > 0 and all(c.replace("-", "").strip() == "" for c in row):
                    continue
                data_rows.append(row)
                
            if len(data_rows) > 0:
                cols_count = len(data_rows[0])
                rows_count = len(data_rows)
                
                table = doc.add_table(rows=rows_count, cols=cols_count)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                # Apply column auto-sizing
                table.autofit = True
                
                for r_idx, row_data in enumerate(data_rows):
                    row = table.rows[r_idx]
                    
                    # Prevent row breaking across pages
                    trPr = row._tr.get_or_add_trPr()
                    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                    
                    # If header row, repeat on every page
                    if r_idx == 0:
                        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                        
                    for c_idx, cell_value in enumerate(row_data):
                        # Ensure index safety
                        if c_idx >= len(row.cells):
                            break
                        cell = row.cells[c_idx]
                        
                        # Set cell margins (padding)
                        set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        
                        if r_idx == 0:
                            # Header Row Styling
                            set_cell_shading(cell, "1E3A8A") # Navy Blue
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(cell_value)
                            run.font.name = 'Calibri'
                            run.font.size = Pt(10)
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255) # White
                        else:
                            # Data Row Styling
                            if r_idx % 2 == 0:
                                set_cell_shading(cell, "F9FAFB") # Zebra striping
                            else:
                                set_cell_shading(cell, "FFFFFF")
                                
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            parse_and_add_runs(p, cell_value)
                            
                # Add spacing after table
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(0)
                spacer.paragraph_format.space_after = Pt(8)
                
            table_rows = []
            # Do not increment idx, process current line again since it's not a table row
            continue
            
        # 3. Handle Regular Elements
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(stripped[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138) # Deep Navy (#1E3A8A)
            
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(stripped[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138) # Deep Navy (#1E3A8A)
            
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(stripped[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(55, 65, 81) # Charcoal Gray
            
        elif stripped.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(stripped[5:])
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(55, 65, 81)
            
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            parse_and_add_runs(p, stripped[2:])
            
        elif stripped.startswith("![") and stripped.endswith(")"):
            # Image pattern: ![caption](path)
            img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
            if img_match:
                caption = img_match.group(1)
                img_path_raw = img_match.group(2)
                # Normalize path
                if img_path_raw.startswith("file:///"):
                    img_path = img_path_raw.replace("file:///", "")
                else:
                    # Resolve relative to md file folder
                    img_path = os.path.join(os.path.dirname(md_path), img_path_raw)
                
                # Check absolute or relative path
                if not os.path.isabs(img_path):
                    img_path = os.path.abspath(img_path)
                
                # Replace Windows forward/back slashes
                img_path = img_path.replace("/", os.sep).replace("\\", os.sep)
                
                if os.path.exists(img_path):
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(8)
                        p_img.paragraph_format.space_after = Pt(4)
                        run_img = p_img.add_run()
                        run_img.add_picture(img_path, width=Inches(5.5))
                        
                        # Add caption
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_before = Pt(0)
                        p_cap.paragraph_format.space_after = Pt(8)
                        run_cap = p_cap.add_run(f"Hình: {caption}")
                        run_cap.font.name = 'Calibri'
                        run_cap.font.size = Pt(9.5)
                        run_cap.italic = True
                        run_cap.font.color.rgb = RGBColor(107, 114, 128)
                    except Exception as img_err:
                        print(f"Error inserting image {img_path}: {img_err}")
                else:
                    print(f"Image file not found: {img_path}")
            
        elif len(stripped) > 0:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            parse_and_add_runs(p, stripped)

            
        else:
            # Empty line, keep small spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            
        idx += 1
        
    try:
        doc.save(docx_path)
        print(f"Successfully generated {os.path.basename(docx_path)}!")
    except PermissionError:
        new_path = docx_path.replace(".docx", "_New.docx")
        doc.save(new_path)
        print(f"WARNING: Permission Denied for {os.path.basename(docx_path)} (locked by Word). Saved as {os.path.basename(new_path)} instead!")

def compile_all():
    doc_dir = "d:/swp/technical_docs"
    sys_docs = [
        "RDS.md",
        "SRS.md",
        "Database_Design.md",
        "Architecture.md",
        "API_Docs.md",
        "Roadmap_Task_Division.md",
        "Tables_And_Fields.md"
    ]
    
    for md_name in sys_docs:
        md_path = os.path.join(doc_dir, md_name)
        docx_name = md_name.replace(".md", ".docx")
        docx_path = os.path.join(doc_dir, docx_name)
        
        if os.path.exists(md_path):
            try:
                compile_markdown_to_docx(md_path, docx_path)
            except Exception as e:
                print(f"ERROR: Could not compile {md_name}: {e}")
        else:
            print(f"Markdown file not found: {md_path}")
            
    print("\n--- ALL DOCUMENTS SUCCESSFULLY COMPILED ---")

if __name__ == "__main__":
    compile_all()
