import os
import docx

def docx_to_md(docx_path, md_path):
    print(f"Converting {docx_path} to {md_path}...")
    doc = docx.Document(docx_path)
    
    with open(md_path, "w", encoding="utf-8") as f:
        # We will iterate through paragraphs and tables in order of their appearance
        # doc.element.body contains all elements in order.
        # We can extract them in order.
        body_elements = doc.element.body
        
        # A map of elements for easy lookups
        paragraphs = {p._element: p for p in doc.paragraphs}
        tables = {t._element: t for t in doc.tables}
        
        for child in body_elements:
            if child in paragraphs:
                p = paragraphs[child]
                text = p.text.strip()
                style_name = p.style.name
                
                if not text:
                    f.write("\n")
                    continue
                
                # Check style for headings
                if style_name.startswith("Heading 1"):
                    f.write(f"\n# {text}\n\n")
                elif style_name.startswith("Heading 2"):
                    f.write(f"\n## {text}\n\n")
                elif style_name.startswith("Heading 3"):
                    f.write(f"\n### {text}\n\n")
                elif style_name.startswith("Heading 4"):
                    f.write(f"\n#### {text}\n\n")
                elif style_name.startswith("Title"):
                    f.write(f"\n# {text}\n\n")
                elif style_name.startswith("Subtitle"):
                    f.write(f"\n### {text}\n\n")
                else:
                    # Normal text, check if it looks like a list
                    if text.startswith("- ") or text.startswith("* "):
                        f.write(f"{text}\n")
                    elif len(text) > 0 and text[0].isdigit() and "." in text[:4] and not text.split(".")[0].isdigit():
                        f.write(f"{text}\n")
                    else:
                        f.write(f"{text}\n\n")
            elif child in tables:
                t = tables[child]
                f.write("\n")
                for row_idx, row in enumerate(t.rows):
                    row_cells = [cell.text.strip().replace("\n", " <br> ") for cell in row.cells]
                    # Format as markdown table
                    f.write("| " + " | ".join(row_cells) + " |\n")
                    if row_idx == 0:
                        # Header separator
                        f.write("| " + " | ".join(["---"] * len(row.cells)) + " |\n")
                f.write("\n")

def convert_all():
    doc_dir = "d:/swp/technical_docs"
    sys_docs = ["API_Docs.docx", "Architecture.docx", "Database_Design.docx", "MultiBranch_Upgrade_Plan.docx", "SRS.docx"]
    for doc_name in sys_docs:
        docx_path = os.path.join(doc_dir, doc_name)
        md_name = doc_name.replace(".docx", ".md")
        md_path = os.path.join(doc_dir, md_name)
        if os.path.exists(docx_path):
            docx_to_md(docx_path, md_path)
            print(f"Successfully converted {doc_name} to {md_name}!")
        else:
            print(f"File not found: {docx_path}")

if __name__ == "__main__":
    convert_all()
